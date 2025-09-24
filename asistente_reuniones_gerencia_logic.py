# asistente_reuniones_gerencia_logic.py (VERSIÓN FINAL Y ROBUSTA)

import pyaudio
import wave
import os
import threading
import speech_recognition as sr
from docx import Document
import time
import json
from concurrent.futures import ThreadPoolExecutor, TimeoutError
import huggingface
from docx.shared import Inches # Importación útil para el futuro si quieres controlar anchos
import pypandoc # <<< Nueva importación


# --- Constantes de Audio (no cambian) ---
CHUNK = 1024
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 16000

# --- Definición de Áreas y Personal Fijo (no cambian) ---
AREAS = {
    "Innovación y Desarrollo": "innovacion_y_desarrollo.txt",
    "Soporte": "soporte.txt",
    "Gestión del Dato": "gestion_del_dato.txt",
    "Gestión de la Información": "gestion_de_la_informacion.txt"
}
JEFE = "Diego mauricio peña Bolaños"
SECRETARIA = "Luz Adriana Ricardo"

# --- Funciones de Lógica ---

def get_integrantes(area_filename):
    """Lee un archivo .txt y devuelve una lista de nombres."""
    try:
        with open(area_filename, 'r', encoding='utf-8') as f:
            return [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        return []

def get_todos_los_integrantes(ruta_carpeta_datos):
    """
    Recorre todas las áreas, lee los archivos de integrantes desde una ruta base
    y devuelve una lista de tuplas (nombre, area).
    """
    todos = []
    for area, filename in AREAS.items():
        ruta_completa = os.path.join(ruta_carpeta_datos, filename)
        integrantes_area = get_integrantes(ruta_completa)
        for nombre in integrantes_area:
            todos.append((nombre, area))
    
    todos.sort(key=lambda x: x[0])
    return todos

def hilo_grabacion_manual(stop_event, frames_buffer):
    """Graba audio en un hilo y lo guarda en un buffer hasta que el evento 'stop' se activa."""
    p = pyaudio.PyAudio()
    stream = p.open(format=FORMAT, channels=CHANNELS, rate=RATE, input=True, frames_per_buffer=CHUNK)
    
    while not stop_event.is_set():
        data = stream.read(CHUNK)
        frames_buffer.append(data)

    stream.stop_stream()
    stream.close()
    p.terminate()

def transcribir_dialogo_aislado(audio_frames):
    """
    Toma frames de audio y devuelve el texto.
    Esta función está diseñada para ser ejecutada en un hilo separado
    y lanzará excepciones en caso de error para que el llamador las maneje.
    """
    if not audio_frames:
        return "[Grabación vacía]"
    
    recognizer = sr.Recognizer()
    audio_data = b''.join(audio_frames)
    audio_file = sr.AudioData(audio_data, RATE, 2) # 2 bytes por muestra para 16-bit

    # Esta función puede lanzar sr.UnknownValueError o sr.RequestError.
    # El hilo que la llama debe estar preparado para capturarlas.
    texto = recognizer.recognize_google(audio_file, language="es-ES")
    return texto

class ActaWord:
    """
    Clase para manejar la creación del documento Word, con un diseño
    robusto que guarda audios primero y permite reanudar transcripciones interrumpidas.
    """
    def __init__(self, titulo, participantes):
        self.titulo = titulo
        self.participantes = participantes
        self.lock = threading.Lock()
        self.cola_de_grabaciones = []

    def agregar_grabacion(self, hablante, audio_data):
        """Añade una grabación a la cola de pendientes en memoria."""
        with self.lock:
            if audio_data:
                self.cola_de_grabaciones.append((hablante, audio_data))

    def _guardar_wav(self, path, audio_data):
        """Función de ayuda para escribir un archivo .wav."""
        p = pyaudio.PyAudio()
        with wave.open(path, 'wb') as wf:
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(RATE)
            wf.writeframes(audio_data)
        p.terminate()

    def guardar_proyecto_para_transcribir(self, ruta_carpeta_reunion):
        """
        PASO 1: Guarda todos los audios en crudo y un archivo de proyecto (.json).
        Este paso es rápido, no necesita internet y asegura los datos.
        """
        try:
            with self.lock:
                os.makedirs(ruta_carpeta_reunion, exist_ok=True)
                
                proyecto_info = {
                    "titulo": self.titulo,
                    "participantes": self.participantes,
                    "dialogos": []
                }

                if not self.cola_de_grabaciones:
                    return None, "No se grabaron diálogos."

                for i, (hablante, audio_data) in enumerate(self.cola_de_grabaciones):
                    nombre_audio = f"dialogo_{i+1}.wav"
                    ruta_audio = os.path.join(ruta_carpeta_reunion, nombre_audio)
                    self._guardar_wav(ruta_audio, audio_data)
                    
                    proyecto_info["dialogos"].append({
                        "id": i + 1,
                        "hablante": hablante,
                        "archivo_audio": nombre_audio,
                        "texto_transcrito": None
                    })

                ruta_proyecto_json = os.path.join(ruta_carpeta_reunion, "proyecto_reunion.json")
                with open(ruta_proyecto_json, 'w', encoding='utf-8') as f:
                    json.dump(proyecto_info, f, indent=4, ensure_ascii=False)
                
                return ruta_proyecto_json, None
        except Exception as e:
            return None, f"Error al guardar los audios del proyecto: {e}"

    def transcribir_desde_proyecto(self, ruta_proyecto_json, update_progress_callback, stop_event=None):
        """
        PASO 2: Lee el proyecto, transcribe los audios con timeouts, y genera el .docx.
        Esta versión es robusta contra bloqueos de red.
        """
        ruta_carpeta_reunion = os.path.dirname(ruta_proyecto_json)
        proyecto_info = {}
        error_encontrado = None

        try:
            with open(ruta_proyecto_json, 'r', encoding='utf-8') as f:
                proyecto_info = json.load(f)

            doc = Document()
            doc.add_heading(proyecto_info["titulo"], level=1)
            doc.add_heading("Acta de Reunión", level=2)
            doc.add_paragraph(f"Fecha: {time.strftime('%d-%m-%Y')}")
            doc.add_paragraph(f"Participantes: {', '.join(proyecto_info['participantes'])}")
            doc.add_paragraph("-" * 50)
            
            total_dialogos = len(proyecto_info["dialogos"])
            
            for i, dialogo in enumerate(proyecto_info["dialogos"]):
                if stop_event and stop_event.is_set():
                    error_encontrado = "Proceso cancelado por pérdida de conexión (detectado por monitor)."
                    break
                
                progreso = (i + 1) / total_dialogos
                hablante = dialogo["hablante"]
                update_progress_callback(progreso, f"Procesando diálogo {i+1}/{total_dialogos} ({hablante})...")
                
                if dialogo.get("texto_transcrito"):
                    texto = dialogo["texto_transcrito"]
                else:
                    try:
                        ruta_audio = os.path.join(ruta_carpeta_reunion, dialogo["archivo_audio"])
                        with wave.open(ruta_audio, 'rb') as wf:
                            frames = wf.readframes(wf.getnframes())
                        
                        # --- INICIO: Lógica de Transcripción con Timeout ---
                        executor = ThreadPoolExecutor(max_workers=1)
                        future = executor.submit(transcribir_dialogo_aislado, [frames])
                        
                        try:
                            # Esperamos el resultado con un timeout de 180 segundos
                            texto = future.result(timeout=180)
                        except TimeoutError:
                            texto = "[Error de Transcripción: La operación tardó demasiado (Timeout)]"
                        except Exception as e:
                            # Capturamos excepciones que ocurrieron DENTRO del hilo de transcripción
                            if isinstance(e, sr.UnknownValueError):
                                texto = "[Audio no reconocido o silencio]"
                            elif isinstance(e, sr.RequestError):
                                texto = f"[Error de Conexión en Transcripción: {e}]"
                            else:
                                texto = f"[Error inesperado durante transcripción: {e}]"
                        # --- FIN: Lógica de Transcripción con Timeout ---
                        
                        dialogo["texto_transcrito"] = texto

                        # Si la transcripción falló por un problema de red, detenemos todo el proceso.
                        if "[Error de Conexión" in texto or "[Timeout]" in texto:
                            error_encontrado = "Se perdió la conexión a internet durante la transcripción."
                            break

                    except Exception as e:
                        error_encontrado = f"Error crítico al leer el archivo de audio para el diálogo {i+1}: {e}"
                        break
                
                p = doc.add_paragraph()
                p.add_run(f"Diálogo {dialogo['id']} - {hablante}: ").bold = True
                p.add_run(texto)

            if error_encontrado:
                raise Exception(error_encontrado) 

            nombre_reunion = os.path.basename(ruta_carpeta_reunion)
            ruta_word = os.path.join(ruta_carpeta_reunion, f"{nombre_reunion}.docx")
            doc.save(ruta_word)
            update_progress_callback(1.0, "¡Acta literal completada!")
            mensaje_exito = f"Acta y audios guardados en la carpeta: {nombre_reunion}"
            return True, mensaje_exito, ruta_word # <--- DEVOLVEMOS LA RUTA

        except Exception as e:
            # --- Bloque de guardado de emergencia ---
            # Guarda el progreso parcial en el archivo JSON antes de salir.
            update_progress_callback(1.0, f"Error: {e}. Guardando progreso...")
            if proyecto_info and "dialogos" in proyecto_info:
                with open(ruta_proyecto_json, 'w', encoding='utf-8') as f:
                    json.dump(proyecto_info, f, indent=4, ensure_ascii=False)
            return False, str(e), None
        

    def generar_acta_inteligente(self, ruta_acta_literal, update_progress_callback):
        """
        PASO 3: Lee un acta literal y la transforma en un acta oficial
        siguiendo la plantilla del HUV.
        """
        update_progress_callback(0.1, "Iniciando formateo de Acta Oficial HUV...")
        try:
            # 1. Leer el texto del acta literal (la transcripción)
            doc_literal = Document(ruta_acta_literal)
            texto_completo_transcripcion = "\n".join([para.text for para in doc_literal.paragraphs if para.text.strip()])
            
            if not texto_completo_transcripcion:
                return False, "La transcripción está vacía, no se puede generar el acta."

            update_progress_callback(0.4, "Transcripción leída. Construyendo prompt para IA...")

            # 2. Construir el prompt de ingeniería avanzada con la plantilla
            #    Nota: Los {{PLACEHOLDERS}} son para datos que la UI debería pedir en el futuro.
            #    Por ahora, la IA los llenará con información genérica que podemos editar.
            
            prompt_plantilla_huv = f"""
Eres un asistente administrativo experto del Hospital Universitario del Valle "Evaristo García" E.S.E., especializado en la redacción de actas de comité. Tu tarea es tomar la transcripción completa de una reunión y la información general, y formatearla en un acta oficial siguiendo ESTRICTAMENTE la siguiente plantilla y estilo.

**PLANTILLA DEL ACTA OFICIAL:**

ACTA No. {{NÚMERO_Y_AÑO_ACTA}}

REUNIÓN DE {{TÍTULO_REUNIÓN}}
DEPARTAMENTO DE {{DEPARTAMENTO}}

1. DATOS GENERALES DE LA REUNIÓN

FECHA: {{FECHA}}
LUGAR: {{LUGAR}}
HORA DE INICIO: {{HORA_INICIO}}     HORA DE FINALIZACIÓN: {{HORA_FIN}}
MODERADOR: {{MODERADOR}}
OBJETIVO DE LA REUNIÓN: {{OBJETIVO}}
ASISTENCIA: Ver listado de asistencia digital.
Verificación del quórum: Sí cumple

ORDEN DEL DÍA
{{LISTA_ORDEN_DÍA}}

DESARROLLO DE LA REUNIÓN
[Aquí debes generar un resumen narrativo de la transcripción. Sigue estas reglas:
- Comienza con una frase de apertura estándar.
- Agrupa las discusiones por tema, no necesariamente en orden cronológico.
- Redacta en tercera persona y tiempo pasado.
- Usa frases conectoras como "Interviene [Nombre y Cargo]:", "[Nombre] informa que...", "[Nombre] responde que...", "Se presenta una propuesta sobre...".
- Extrae los puntos clave, problemas, soluciones y decisiones. Omite el relleno.]

REVISIÓN Y APROBACIÓN DEL ACTA:
NOMBRE                  CARGO                            FIRMA
{{NOMBRE_REVISOR_1}}    {{CARGO_REVISOR_1}}
{{NOMBRE_REVISOR_2}}    {{CARGO_REVISOR_2}}

REVISIÓN DE COMPROMISOS PREVIOS:
COMPROMISO | RESPONSABLE | FECHA DE CUMPLIMIENTO | ESTADO
{{LISTA_COMPROMISOS_PREVIOS}}

COMPROMISOS:
[Analiza la sección 'DESARROLLO DE LA REUNIÓN' que generaste e identifica cualquier tarea, compromiso o acción nueva que se haya asignado. Formatea como una lista o tabla con COMPROMISO, RESPONSABLE y FECHA DE CUMPLIMIENTO si se menciona. Si no hay compromisos nuevos, escribe "No se generaron nuevos compromisos en esta reunión."]

Elaboro: {{NOMBRE_ELABORADOR}}
Reviso: {{NOMBRE_REVISOR_JEFE}}

**INFORMACIÓN PROPORCIONADA PARA EL ACTA:**

- **Transcripción Completa:**
---
{texto_completo_transcripcion}
---
- **Título de la Reunión:** {self.titulo}
- **Participantes:** {', '.join(self.participantes)}
- **Fecha:** {time.strftime('%d/%m/%Y')}
- **Elaborador:** Luz Adriana Ricardo
- **Revisor Jefe:** Diego Mauricio Peña Bolaños

**ACCIÓN REQUERIDA:**
Ahora, genera el contenido completo del acta final rellenando la plantilla anterior con la información proporcionada y el análisis de la transcripción.
"""

            # 3. Llamar a la IA con el nuevo prompt
            update_progress_callback(0.5, "Contactando a la IA para generar el acta formateada...")
            acta_formateada = huggingface.resumir_intervencion_hf(
                texto_literal=prompt_plantilla_huv,
                hablante="Generador de Acta HUV",
                participantes=[]
            )

            # ... (código anterior para obtener `acta_formateada` de la IA) ...

            update_progress_callback(0.8, "Acta en Markdown recibida. Convirtiendo a formato Word con Pandoc...")

            # --- INICIO DE LA MODIFICACIÓN CON PYPANDOC ---
            
            # 4. Convertir el texto Markdown a un archivo .docx usando pypandoc.
            
            # Renombramos la variable para mayor claridad.
            acta_en_markdown = acta_formateada
            
            # Definimos la ruta de salida, igual que antes.
            ruta_carpeta = os.path.dirname(ruta_acta_literal)
            titulo_original = os.path.basename(ruta_acta_literal).replace('.docx', '')
            nombre_acta_final = f"{titulo_original}_ActaOficial_HUV.docx"
            ruta_acta_final = os.path.join(ruta_carpeta, nombre_acta_final)
            
            try:
                # Esta es la línea que hace toda la magia con Pandoc.
                # Le decimos que convierta el texto `acta_en_markdown` al formato 'docx',
                # asumiendo que el formato de entrada es 'md' (markdown),
                # y que guarde el resultado directamente en `ruta_acta_final`.
                pypandoc.convert_text(
                    acta_en_markdown, 
                    'docx', 
                    format='md', 
                    outputfile=ruta_acta_final,
                    # Opcional: añade argumentos extra si necesitas una plantilla de Word específica
                    # extra_args=['--reference-doc=mi_plantilla.docx'] 
                )

            except OSError:
                # Este bloque se ejecutará si pypandoc no puede encontrar a Pandoc.
                error_msg = ("Error Crítico: No se pudo encontrar Pandoc. "
                             "Asegúrate de que esté instalado en el sistema y disponible en el PATH. "
                             "No se pudo generar el acta formateada.")
                # Devolvemos False para que la UI pueda mostrar el error.
                return False, error_msg

            # --- FIN DE LA MODIFICACIÓN ---

            update_progress_callback(1.0, "¡Acta Oficial HUV generada con éxito!")

            return True, f"Acta Oficial guardada como {nombre_acta_final}"

        except Exception as e:
            # Este bloque ahora capturará otros posibles errores, 
            # como problemas al escribir el archivo, etc.
            return False, f"Error al generar el Acta Oficial: {e}"